import os
import csv
import chardet
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from translator import Translator


def detect_encoding(file_path):
    """Detects the encoding of a file.

    Args:
        file_path (str): The path to the file to detect the encoding of.

    Returns:
        str: The detected encoding of the file.
    """
    with open(file_path, 'rb') as f:
        raw_data = f.read(10000)
    result = chardet.detect(raw_data)
    return result['encoding']


def process_word_document(csv_filepath, template_filepath, output_filepath):
    
    """
    Processes a CSV file containing titles and fields to translate and adds them
    to a Word document template, saving the translated document to a new file.

    Args:
        csv_filepath (str): The path to the CSV file to process.
        template_filepath (str): The path to the Word document template file.
        output_filepath (str): The path where the translated document should be saved.

    The function detects the encoding of the CSV file, loads it, and translates
    each title and field using a Translator object. It adds the translated titles
    as H2 headings and fields as paragraphs to the Word document template, and
    saves the translated document to the specified output file path.
    """
    encoding = detect_encoding(csv_filepath)

    with open(csv_filepath, 'r', encoding=encoding) as csv_file:
        reader = csv.DictReader(csv_file)
        data = [row for row in reader]

    translator = Translator()

    doc = Document(template_filepath)

    for row in data:

        translated_title = translator.translate(row['Title'])
        doc.add_heading(translated_title, level=2)

        for key, value in row.items():
            if key != 'Title':
                translated_key = translator.translate(key)
                translated_value = translator.translate(value)
            
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
                run = paragraph.add_run(translated_key)
                font = run.font
                font.bold = True
            
                run = paragraph.add_run(":")
                paragraph.add_run().add_break()
            
                run = paragraph.add_run(translated_value)
                font = run.font
                if key == 'RiskRating':
                    font.bold = True
                
                    if translated_value == 'Crítico':
                        font.color.rgb = RGBColor(139, 0, 0)  # Dark red
                    elif translated_value == 'Alto':
                        font.color.rgb = RGBColor(255, 0, 0)  # Red
                    elif translated_value == 'Medio':
                        font.color.rgb = RGBColor(222, 169, 0)  # Gold
                    elif translated_value == 'Bajo':
                        font.color.rgb = RGBColor(0, 128, 0)  # Green
                    elif translated_value == 'Informativo':
                        font.color.rgb = RGBColor(0, 0, 255)  # Blue

    doc.save(output_filepath)


if __name__ == '__main__':
    csv_filepath = os.path.join(os.path.dirname(__file__), "./combined_output.csv")
    template_filepath = os.path.join(os.path.dirname(__file__), "./template.docx")
    output_filepath = os.path.join(os.path.dirname(__file__), "[Nombre de cliente] ASSESMENT - nombre del tenant.docx")
    
    print("Word translation started")

    process_word_document(csv_filepath, template_filepath, output_filepath)

    print("Word translation finished")
