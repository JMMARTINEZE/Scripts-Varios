import os
import csv
import chardet
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from translator import Translator
import pandas as pd
import logging
from dotenv import load_dotenv

load_dotenv()

input_html_folder = os.getenv('INPUT_HTML_FOLDER')
output_folder = os.getenv('OUTPUT_FOLDER')

def process_word_document(csv_filepath, template_filepath, output_folder):
    try:
        # Detectar la codificación del archivo CSV
        encoding = detect_encoding(csv_filepath)
        if encoding is None:
            logging.error("Error detecting encoding")
            return

        # Leer el archivo CSV con la codificación detectada
        with open(csv_filepath, 'r', encoding=encoding) as csv_file:
            reader = csv.DictReader(csv_file)
            data = [row for row in reader]

        # Traductor
        translator = Translator()

        # Cargar el documento Word existente como plantilla
        doc = Document(template_filepath)

        # Traducir y agregar los campos
        translated_data = []
        for row in data:
            translated_row = {}
            for key, value in row.items():
                try:
                    translated_key = translator.translate(key)
                    translated_value = translator.translate(value)
                    translated_row[translated_key] = translated_value
                    # Agregar los campos al documento Word
                    if key != 'Title':
                        paragraph = doc.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run = paragraph.add_run(translated_key)
                        font = run.font
                        font.bold = True
                        run = paragraph.add_run(":")
                        paragraph.add_run().add_break()
                        run = paragraph.add_run(translated_value)
                        font = run.font
                        if key == 'RiskRating':
                            font.bold = True
                            # Set font color
                            if translated_value == 'Crítico':
                                font.color.rgb = RGBColor(139, 0, 0)  # Dark red
                            elif translated_value == 'Alto':
                                font.color.rgb = RGBColor(255, 0, 0)  # Red
                            elif translated_value == 'Medio':
                                font.color.rgb = RGBColor(222, 169, 0)  # Gold
                            elif translated_value == 'Bajo':
                                font.color.rgb = RGBColor(0, 128, 0)  # Green
                            elif translated_value == 'Informativo':
                                font.color.rgb = RGBColor(0, 0, 255)  # Blue
                except Exception as e:
                    logging.error(f"Error translating row: {e}")
            translated_data.append(translated_row)

        # Crear un DataFrame con el texto traducido
        df = pd.DataFrame(translated_data)

        # Escribir el DataFrame en un archivo CSV
        html_filename = os.path.basename(csv_filepath).replace('.csv', '.html')
        output_filename = f"{html_filename}.docx"
        output_csv_filename = f"{html_filename}.csv"
        df.to_csv(os.path.join(output_folder, output_csv_filename), index=False)

        # Guardar el documento respetando la plantilla original
        doc.save(os.path.join(output_folder, output_filename))
    except Exception as e:
        logging.error(f"Error processing word document: {e}")

if __name__ == '__main__':
    csv_filepath = os.path.join(input_html_folder, "combined_output.csv")
    template_filepath = os.path.join(input_html_folder, "template.docx")
    process_word_document(csv_filepath, template_filepath, output_folder)
    print("Word translation finished")