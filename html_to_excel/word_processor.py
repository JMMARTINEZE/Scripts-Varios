import os
import csv
import chardet
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from translator import Translator
import pandas as pd


def detect_encoding(file_path):
    with open(file_path, 'rb') as f:
        raw_data = f.read(10000)
    result = chardet.detect(raw_data)
    return result['encoding']


def process_word_document(csv_filepath, template_filepath, output_filepath, output_csv_filepath):
    # Detectar la codificación del archivo CSV
    encoding = detect_encoding(csv_filepath)

    # Leer el archivo CSV con la codificación detectada
    with open(csv_filepath, 'r', encoding=encoding) as csv_file:
        reader = csv.DictReader(csv_file)
        data = [row for row in reader]

    # Traductor
    translator = Translator()

    # Cargar el documento Word existente como plantilla
    doc = Document(template_filepath)

    # Traducir y agregar los campos
    translated_data = []
    for row in data:
        translated_row = {}
        for key, value in row.items():
            translated_key = translator.translate(key)
            translated_value = translator.translate(value)
            translated_row[translated_key] = translated_value
            # Agregar los campos al documento Word
            if key != 'Title':
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = paragraph.add_run(translated_key)
                font = run.font
                font.bold = True
                run = paragraph.add_run(":")
                paragraph.add_run().add_break()
                run = paragraph.add_run(translated_value)
                font = run.font
                if key == 'RiskRating':
                    font.bold = True
                    # Set font color
                    if translated_value == 'Crítico':
                        font.color.rgb = RGBColor(139, 0, 0)  # Dark red
                    elif translated_value == 'Alto':
                        font.color.rgb = RGBColor(255, 0, 0)  # Red
                    elif translated_value == 'Medio':
                        font.color.rgb = RGBColor(222, 169, 0)  # Gold
                    elif translated_value == 'Bajo':
                        font.color.rgb = RGBColor(0, 128, 0)  # Green
                    elif translated_value == 'Informativo':
                        font.color.rgb = RGBColor(0, 0, 255)  # Blue
        translated_data.append(translated_row)

    # Crear un DataFrame con el texto traducido
    df = pd.DataFrame(translated_data)

    # Escribir el DataFrame en un archivo CSV
    df.to_csv(output_csv_filepath, index=False)

    # Guardar el documento respetando la plantilla original
    doc.save(output_filepath)


if __name__ == '__main__':
    csv_filepath = os.path.join(os.path.dirname(__file__), "./combined_output.csv")
    template_filepath = os.path.join(os.path.dirname(__file__), "./template.docx")
    output_filepath = os.path.join(os.path.dirname(__file__), "traducciones_respetando_template.docx")
    output_csv_filepath = os.path.join(os.path.dirname(__file__), "translated_text.csv")

    print("Word translation started")

    process_word_document(csv_filepath, template_filepath, output_filepath, output_csv_filepath)

    print("Word translation finished")