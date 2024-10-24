import os
import csv
import chardet
from deep_translator import GoogleTranslator
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Detectar la codificación del archivo
def detect_encoding(file_path):
    with open(file_path, 'rb') as f:
        raw_data = f.read(10000)  # Leer los primeros 10,000 bytes
    result = chardet.detect(raw_data)
    return result['encoding']

# Ruta del archivo CSV y la plantilla de Word
csv_filepath = os.path.join(os.path.dirname(__file__), "./combined_output.csv")
template_filepath = os.path.join(os.path.dirname(__file__), "./template.docx")

# Detectar la codificación del archivo CSV
encoding = detect_encoding(csv_filepath)
print(f"Codificación detectada: {encoding}")

# Leer el archivo CSV con la codificación detectada
with open(csv_filepath, 'r', encoding=encoding) as csv_file:
    reader = csv.DictReader(csv_file)
    data = [row for row in reader]

# Traductor
translator = GoogleTranslator(source='auto', target='es')

# Cargar el documento Word existente como plantilla
doc = Document(template_filepath)

# Traducir y agregar contenido al documento
for row in data:
    # Traducir el título y agregar como H2
    translated_title = translator.translate(row['Title'])
    doc.add_heading(translated_title, level=2)

    # Traducir y agregar los campos
    for key, value in row.items():
        if key != 'Title':
            translated_key = translator.translate(key)
            translated_value = translator.translate(value)
    
            # Agregar el campo en negrita
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(translated_key)
            font = run.font
            font.bold = True
    
            # Agregar los puntos y el salto de línea
            run = paragraph.add_run(":")
            paragraph.add_run().add_break()
    
            # Agregar el contenido del campo
            paragraph.add_run(translated_value)

# Guardar el documento respetando la plantilla original
output_filepath = os.path.join(os.path.dirname(__file__), "traducciones_respetando_plantilla.docx")
doc.save(output_filepath)

print(f"Documento guardado en: {output_filepath}")